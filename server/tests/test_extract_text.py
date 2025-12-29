"""Unit tests for text extraction helpers."""
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from app.services.extractor import (
    ExtractionError,
    extract,
    extract_docx,
    extract_pdf,
    extract_txt_or_md,
)


def test_extract_txt_or_md_reads_file(tmp_path: Path) -> None:
    # Basic text extraction should preserve content.
    path = tmp_path / "note.md"
    path.write_text("hello world", encoding="utf-8")

    assert extract_txt_or_md(path) == "hello world"


def test_extract_docx_reads_paragraphs(tmp_path: Path) -> None:
    # Docx paragraphs should be joined with newlines.
    path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.save(str(path))

    assert extract_docx(path) == "alpha\nbeta"


def test_extract_pdf_raises_on_short_text(tmp_path: Path) -> None:
    # Blank PDF should fail the min-chars guard.
    path = tmp_path / "short.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as handle:
        writer.write(handle)

    with pytest.raises(ExtractionError):
        extract_pdf(path)


def test_extract_dispatches_by_extension(tmp_path: Path) -> None:
    # Entry-point should dispatch by extension and return meta when present.
    txt_path = tmp_path / "note.txt"
    txt_path.write_text("hello", encoding="utf-8")
    text, meta = extract(str(txt_path), "note.txt", "text/plain")
    assert text == "hello"
    assert meta is None

    docx_path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("hello docx")
    doc.save(str(docx_path))
    text, meta = extract(str(docx_path), "note.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert text == "hello docx"
    assert meta is None


def test_extract_raises_on_missing_file(tmp_path: Path) -> None:
    # Missing files should raise a controlled error.
    missing = tmp_path / "missing.txt"
    with pytest.raises(ExtractionError):
        extract(str(missing), "missing.txt", "text/plain")


def test_extract_raises_on_unsupported_extension(tmp_path: Path) -> None:
    # Unsupported extensions should raise a controlled error.
    path = tmp_path / "note.bin"
    path.write_bytes(b"data")
    with pytest.raises(ExtractionError):
        extract(str(path), "note.bin", "application/octet-stream")
